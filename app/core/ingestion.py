import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

import logging
from typing import List, Optional
from datetime import datetime
import subprocess
import json
from ..models import DocumentMetadata, DocumentChunk, IngestedDocument
from pdfminer.high_level import extract_text
from pdfminer.pdfparser import PDFSyntaxError
from docx import Document as DocxDocument
import pandas as pd
from .chunking import chunk_text, ChunkStrategy

try:
    print("DEBUG: Attempting to import markitdown...")
    from markitdown import MarkItDown
    MARKITDOWN_AVAILABLE = True
    print("DEBUG: markitdown import succeeded.")
except ImportError as e:
    MARKITDOWN_AVAILABLE = False
    print(f"DEBUG: markitdown import failed: {e}")

# --- MarkItDown Extraction (with chunking options) ---
def extract_with_markitdown(
    path: Path,
    chunk_strategy: ChunkStrategy = "char",
    chunk_size: int = 512,
    overlap: int = 50
) -> IngestedDocument:
    if not MARKITDOWN_AVAILABLE:
        raise ImportError("markitdown is not installed. Please install with pip install 'markitdown[all]'")
    md = MarkItDown(enable_plugins=False)
    try:
        result = md.convert(str(path))
        text = result.text_content
        meta = DocumentMetadata(
            filename=path.name,
            filetype=path.suffix.lstrip('.').lower(),
            extra={"parser": "markitdown"}
        )
        # --- Chunking ---
        chunks = chunk_text(text, strategy=chunk_strategy, chunk_size=chunk_size, overlap=overlap)
        return IngestedDocument(metadata=meta, chunks=chunks)
    except Exception as e:
        ext = path.suffix.lower()
        if ext == '.txt':
            logging.warning(f"[MarkItDown failed for TXT: {e}] Falling back to UTF-8 native extractor.")
            return extract_txt(path)
        elif ext == '.docx':
            logging.warning(f"[MarkItDown failed for DOCX: {e}] Falling back to python-docx extractor.")
            return extract_docx(path)
        elif ext == '.csv':
            logging.warning(f"[MarkItDown failed for CSV: {e}] Falling back to pandas extractor.")
            return extract_csv(path)
        else:
            logging.error(f"[MarkItDown failed for {ext.upper()}: {e}]. No fallback available.")
            raise

# --- TXT and Markdown Extraction ---
def extract_txt(path: Path) -> IngestedDocument:
    text = path.read_text(encoding='utf-8')
    meta = DocumentMetadata(
        filename=path.name,
        filetype='txt',
        extra={}
    )
    chunk = DocumentChunk(text=text)
    return IngestedDocument(metadata=meta, chunks=[chunk])

def extract_markdown(path: Path) -> IngestedDocument:
    # Try markitdown (Node.js) first
    try:
        result = subprocess.run([
            'npx', 'markitdown', '--ast', str(path)
        ], capture_output=True, text=True, check=True)
        ast = json.loads(result.stdout)
        # For now, flatten AST to plain text (structure-aware chunking later)
        text = flatten_markdown_ast(ast)
        used = 'markitdown'
    except Exception:
        # Fallback to markdown-it-py
        try:
            import markdown_it
            md = markdown_it.MarkdownIt()
            text = md.render(path.read_text(encoding='utf-8'))
            used = 'markdown-it-py'
        except ImportError:
            # Fallback to mistune
            import mistune
            text = mistune.create_markdown()(path.read_text(encoding='utf-8'))
            used = 'mistune'
    meta = DocumentMetadata(
        filename=path.name,
        filetype='md',
        extra={'parser': used}
    )
    chunk = DocumentChunk(text=text)
    return IngestedDocument(metadata=meta, chunks=[chunk])

def flatten_markdown_ast(ast: dict) -> str:
    # Simple recursive flatten for MVP
    if isinstance(ast, dict):
        children = ast.get('children', [])
        if 'value' in ast:
            return ast['value'] + '\n' + '\n'.join(flatten_markdown_ast(child) for child in children)
        else:
            return '\n'.join(flatten_markdown_ast(child) for child in children)
    elif isinstance(ast, list):
        return '\n'.join(flatten_markdown_ast(item) for item in ast)
    else:
        return str(ast)

# --- DOCX Fallback ---
def extract_docx(path: Path) -> IngestedDocument:
    doc = DocxDocument(str(path))
    text = '\n'.join([para.text for para in doc.paragraphs])
    meta = DocumentMetadata(
        filename=path.name,
        filetype='docx',
        extra={"parser": "python-docx"}
    )
    chunk = DocumentChunk(text=text)
    return IngestedDocument(metadata=meta, chunks=[chunk])

# --- CSV Fallback ---
def extract_csv(path: Path) -> IngestedDocument:
    df = pd.read_csv(str(path), encoding='utf-8', dtype=str)
    text = df.to_csv(index=False)
    meta = DocumentMetadata(
        filename=path.name,
        filetype='csv',
        extra={"parser": "pandas"}
    )
    chunk = DocumentChunk(text=text)
    return IngestedDocument(metadata=meta, chunks=[chunk])

# --- PDF Extraction ---
def extract_pdf(path: Path) -> IngestedDocument:
    try:
        text = extract_text(str(path))
    except PDFSyntaxError:
        text = "[PDF parsing error]"
    meta = DocumentMetadata(
        filename=path.name,
        filetype='pdf',
        extra={}
    )
    chunk = DocumentChunk(text=text)
    return IngestedDocument(metadata=meta, chunks=[chunk])

# --- CLI Test ---
if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("file", help="File to ingest")
    parser.add_argument("--chunk-strategy", choices=["char", "word", "paragraph"], default="char")
    parser.add_argument("--chunk-size", type=int, default=512)
    parser.add_argument("--overlap", type=int, default=50)
    args = parser.parse_args()
    doc = extract_with_markitdown(
        Path(args.file),
        chunk_strategy=args.chunk_strategy,
        chunk_size=args.chunk_size,
        overlap=args.overlap
    )
    print(f"{len(doc.chunks)} chunks extracted.")
    for i, chunk in enumerate(doc.chunks[:3]):
        print(f"Chunk {i}: {chunk.text[:100].replace('\n', ' ')}{'...' if len(chunk.text) > 100 else ''}")

    # PATCH: Embed and ingest chunks into ChromaDB
    from app.core.embedding import embed_document
    from app.vectorstore import add_chunk_embeddings
    chunk_embs = embed_document(doc)
    add_chunk_embeddings(chunk_embs, doc.metadata.filename)
    print(f"Ingested {len(chunk_embs)} chunks into ChromaDB for document {doc.metadata.filename}")
